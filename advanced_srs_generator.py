from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import io

from diagram_generator import (
    generate_requirements_pie,
    generate_use_case_diagram,
    generate_architecture_diagram,
    generate_dfd,
    generate_nfr_bar,
    generate_user_flow,
)

# ── Helpers ───────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color_hex="1F4E79"):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16)
        )
    return h

def add_diagram(doc, image_bytes: bytes, caption: str, width_inches=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Inches(width_inches))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_paragraph()

def add_req_table(doc, reqs, headers, col_widths, header_bg="2E75B6"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, (cell, text, w) in enumerate(zip(hdr_cells, headers, col_widths)):
        cell.width = Inches(w)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # Data rows
    row_colors = ["FFFFFF", "EBF3FB"]
    for ri, req in enumerate(reqs):
        row_cells = table.add_row().cells
        bg = row_colors[ri % 2]
        values = [req.get(k, "") for k in ["id", "title", "description"]]
        if "category" in req and len(headers) == 4:
            values = [req.get("id",""), req.get("category",""),
                      req.get("title",""), req.get("description","")]
        for ci, (cell, val, w) in enumerate(zip(row_cells, values, col_widths)):
            cell.width = Inches(w)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()


# ── Main Generator ────────────────────────────────────────────
def generate_advanced_srs(project_name: str, requirements: dict,
                           project_desc: str = "") -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    functional     = requirements.get("functional", [])
    non_functional = requirements.get("non_functional", [])

    # ── Cover Page ────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("SOFTWARE REQUIREMENTS\nSPECIFICATION")
    tr.bold = True
    tr.font.size = Pt(28)
    tr.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()
    proj_p = doc.add_paragraph()
    proj_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = proj_p.add_run(project_name)
    pr.bold = True
    pr.font.size = Pt(20)
    pr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_p.add_run(
        f"Generated: {datetime.now().strftime('%B %d, %Y')}\n"
        f"Version: 1.0  |  Classification: Internal\n"
        f"Total Requirements: {len(functional) + len(non_functional)}"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ── TOC placeholder ───────────────────────────────────────
    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Introduction",
        "2. Project Overview & Requirements Summary",
        "3. Use Case Diagram",
        "4. System Architecture",
        "5. Data Flow Diagram",
        "6. Functional Requirements",
        "7. Non-Functional Requirements",
        "8. User Activity Flow",
        "9. Assumptions & Constraints",
        "10. Glossary",
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Inches(0.3)
        p.runs[0].font.size = Pt(11)
    doc.add_page_break()

    # ── 1. Introduction ───────────────────────────────────────
    add_heading(doc, "1. Introduction", level=1)
    add_heading(doc, "1.1 Purpose", level=2, color_hex="2E75B6")
    doc.add_paragraph(
        "This Software Requirements Specification (SRS) document provides a complete "
        "description of all functionality and constraints for the system. It is intended "
        "for project stakeholders, developers, testers, and project managers."
    )
    add_heading(doc, "1.2 Project Scope", level=2, color_hex="2E75B6")
    scope_text = project_desc if project_desc else (
        f"The '{project_name}' system is designed to deliver a robust, scalable, and "
        "user-friendly solution that meets the identified business requirements."
    )
    doc.add_paragraph(scope_text)

    add_heading(doc, "1.3 Document Conventions", level=2, color_hex="2E75B6")
    conv_table = doc.add_table(rows=3, cols=2)
    conv_table.style = "Table Grid"
    for r, (term, meaning) in enumerate([
        ("FR-XXX", "Functional Requirement identifier"),
        ("NFR-XXX", "Non-Functional Requirement identifier"),
        ("TBD", "To Be Determined"),
    ]):
        cells = conv_table.rows[r].cells
        set_cell_bg(cells[0], "EBF3FB")
        cells[0].paragraphs[0].add_run(term).bold = True
        cells[1].paragraphs[0].add_run(meaning)
    doc.add_paragraph()
    doc.add_page_break()

    # ── 2. Requirements Summary + Pie Chart ──────────────────
    add_heading(doc, "2. Project Overview & Requirements Summary", level=1)
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = "Table Grid"
    for r, (label, value) in enumerate([
        ("Project Name", project_name),
        ("Functional Requirements", str(len(functional))),
        ("Non-Functional Requirements", str(len(non_functional))),
        ("Total Requirements", str(len(functional) + len(non_functional))),
    ]):
        cells = summary_table.rows[r].cells
        set_cell_bg(cells[0], "2E75B6")
        lr = cells[0].paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        vr = cells[1].paragraphs[0].add_run(value)
        vr.bold = True
        vr.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    doc.add_paragraph()

    # Pie chart
    pie_bytes = generate_requirements_pie(functional, non_functional)
    add_diagram(doc, pie_bytes, "Figure 1: Requirements Distribution by Category", 5.5)
    doc.add_page_break()

    # ── 3. Use Case Diagram ───────────────────────────────────
    add_heading(doc, "3. Use Case Diagram", level=1)
    doc.add_paragraph(
        "The use case diagram below illustrates the interactions between users "
        "and the system, showing the primary functional capabilities."
    )
    uc_bytes = generate_use_case_diagram(project_name, functional)
    add_diagram(doc, uc_bytes, "Figure 2: System Use Case Diagram", 6.5)
    doc.add_page_break()

    # ── 4. System Architecture ────────────────────────────────
    add_heading(doc, "4. System Architecture Diagram", level=1)
    doc.add_paragraph(
        "The three-tier architecture separates the Presentation, Application, "
        "and Data layers, ensuring scalability, maintainability, and separation of concerns."
    )
    arch_bytes = generate_architecture_diagram(project_name)
    add_diagram(doc, arch_bytes, "Figure 3: Three-Tier System Architecture", 6.5)
    doc.add_page_break()

    # ── 5. DFD ───────────────────────────────────────────────
    add_heading(doc, "5. Data Flow Diagram (Level 0)", level=1)
    doc.add_paragraph(
        "The Data Flow Diagram (DFD) at Level 0 shows the high-level flow of "
        "information between external entities, the system, and data stores."
    )
    dfd_bytes = generate_dfd(project_name, functional)
    add_diagram(doc, dfd_bytes, "Figure 4: Data Flow Diagram — Level 0", 6.5)
    doc.add_page_break()

    # ── 6. Functional Requirements ────────────────────────────
    add_heading(doc, "6. Functional Requirements", level=1)
    doc.add_paragraph(
        "Functional requirements define the specific behaviors, features, and "
        "functions the system must support."
    )
    if functional:
        add_req_table(
            doc, functional,
            ["ID", "Title", "Description"],
            [0.8, 1.8, 4.2],
            header_bg="2E75B6"
        )
    else:
        doc.add_paragraph("No functional requirements identified.")
    doc.add_page_break()

    # ── 7. Non-Functional Requirements ───────────────────────
    add_heading(doc, "7. Non-Functional Requirements", level=1)
    doc.add_paragraph(
        "Non-functional requirements define the quality attributes, constraints, "
        "and system properties the system must satisfy."
    )

    # NFR bar chart
    nfr_bytes = generate_nfr_bar(non_functional)
    add_diagram(doc, nfr_bytes, "Figure 5: NFR Distribution by Category", 5.5)

    # Group NFRs by category
    if non_functional:
        categories = {}
        for req in non_functional:
            cat = req.get("category", "General")
            categories.setdefault(cat, []).append(req)

        for cat, reqs in categories.items():
            add_heading(doc, f"7.x {cat} Requirements", level=2, color_hex="2E75B6")
            add_req_table(
                doc, reqs,
                ["ID", "Category", "Title", "Description"],
                [0.7, 1.2, 1.7, 3.2],
                header_bg="70AD47"
            )
    else:
        doc.add_paragraph("No non-functional requirements identified.")
    doc.add_page_break()

    # ── 8. User Activity Flow ─────────────────────────────────
    add_heading(doc, "8. User Activity Flow Diagram", level=1)
    doc.add_paragraph(
        "The activity diagram shows the sequential flow of user interactions "
        "through the primary functional requirements of the system."
    )
    flow_bytes = generate_user_flow(functional)
    add_diagram(doc, flow_bytes, "Figure 6: User Activity Flow Diagram", 5.0)
    doc.add_page_break()

    # ── 9. Assumptions & Constraints ─────────────────────────
    add_heading(doc, "9. Assumptions & Constraints", level=1)
    add_heading(doc, "9.1 Assumptions", level=2, color_hex="2E75B6")
    assumptions = [
        "Users have access to a stable internet connection.",
        "The system will be hosted on a cloud-based infrastructure.",
        "All users will have valid authentication credentials.",
        "Third-party APIs and services will maintain 99%+ uptime.",
    ]
    for a in assumptions:
        p = doc.add_paragraph(a, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    add_heading(doc, "9.2 Constraints", level=2, color_hex="2E75B6")
    constraints = [
        "The system must comply with data protection regulations (GDPR/local laws).",
        "Development timeline is fixed per the project schedule.",
        "Budget constraints limit third-party service integrations.",
        "The system must support modern browsers (Chrome, Firefox, Edge, Safari).",
    ]
    for c in constraints:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── 10. Glossary ──────────────────────────────────────────
    add_heading(doc, "10. Glossary", level=1)
    glossary = [
        ("SRS", "Software Requirements Specification"),
        ("FR",  "Functional Requirement"),
        ("NFR", "Non-Functional Requirement"),
        ("DFD", "Data Flow Diagram"),
        ("API", "Application Programming Interface"),
        ("UI",  "User Interface"),
        ("DB",  "Database"),
        ("SSL", "Secure Sockets Layer"),
    ]
    g_table = doc.add_table(rows=1, cols=2)
    g_table.style = "Table Grid"
    hdr = g_table.rows[0].cells
    for i, (cell, h) in enumerate(zip(hdr, ["Term", "Definition"])):
        set_cell_bg(cell, "1F4E79")
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for ri, (term, defn) in enumerate(glossary):
        row = g_table.add_row().cells
        bg = "FFFFFF" if ri % 2 == 0 else "EBF3FB"
        set_cell_bg(row[0], bg)
        set_cell_bg(row[1], bg)
        row[0].paragraphs[0].add_run(term).bold = True
        row[1].paragraphs[0].add_run(defn)

    # ── Save ──────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()